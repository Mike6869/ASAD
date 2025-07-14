from docx import Document
from docx.text.paragraph import Paragraph
from docx.document import Document as Document_class
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from collections import namedtuple
import re
import warnings
from docx.opc.exceptions import PackageNotFoundError
from config import ABBREVIATION_TABLE_TEMPLATE, GLOSSARY_PARAGRAPH_TITLES

AbbreviationTuple = namedtuple("AbbreviationTuple", ["abbreviation", "description"])
re_spaces = re.compile(r'\s+')
warnings.filterwarnings("ignore", category=UserWarning, module='docx')


def get_table_header(table):
    return [col.text for col in table.rows[0].cells]


def get_table_rows(table, rows=0, offset=1):
    if rows == 0 or len(table.rows) < rows:
        rows_return_count = len(table.rows)
    else:
        rows_return_count = rows

    returning_rows = []
    for i in range(offset, rows_return_count):
        cur_row = [cell.text for cell in table.rows[i].cells]
        returning_rows.append(cur_row)

    return returning_rows


def find_terms(header):
    # Returns [term_ind, desc_ind] if the table has terms. Otherwise, returns None
    term_ind = None
    desc_ind = None
    for ind, col_name in enumerate(header):
        col_name = re_spaces.sub(' ', col_name.strip().upper())
        if col_name in ABBREVIATION_TABLE_TEMPLATE["term"]:
            term_ind = ind
            continue
        if col_name in ABBREVIATION_TABLE_TEMPLATE["description"]:
            desc_ind = ind

    if (term_ind is not None) and (desc_ind is not None):
        return [term_ind, desc_ind]


def parse_tables(doc_path):
    doc = Document(doc_path)
    table_list = doc.tables
    parsed_tables = []
    for i in range(len(table_list)):
        table = table_list[i]
        table_header = get_table_header(table)
        table_rows = get_table_rows(table)
        table_info = {'header': table_header, 'rows': table_rows}
        parsed_tables.append(table_info)
    return parsed_tables


def get_cell_text(cell):
    """Получение текста ячейки с защитой от ошибок"""
    try:
        return cell.text.strip() if hasattr(cell, 'text') and cell.text else ''
    except:
        return ''


def iter_block_items(parent):
    if isinstance(parent, Document_class):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parsing doc blocks went wrong")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_abbreviations(file):
    doc = Document(file)
    parsed_abbrs = []
    prev_block = None
    
    for block in iter_block_items(doc):
        if not isinstance(block, Table):
            prev_block = block
            continue

        try:
            table_header = get_table_header(block)
            terms_idx = find_terms(table_header)
            
            if terms_idx:
                for row in get_table_rows(block):
                    # Проверяем, что индексы существуют в row
                    if len(row) > max(terms_idx):
                        parsed_abbrs.append(AbbreviationTuple(
                            abbreviation=row[terms_idx[0]], 
                            description=row[terms_idx[1]]
                        ))
            
            # Обработка глоссария без заголовков
            elif isinstance(prev_block, Paragraph) and prev_block.text in GLOSSARY_PARAGRAPH_TITLES:
                for row in get_table_rows(block, offset=0):
                    # Проверяем, что в строке достаточно элементов
                    if len(row) >= 3:
                        parsed_abbrs.append(AbbreviationTuple(
                            abbreviation=row[0], 
                            description=row[2]
                        ))
                        
        except Exception as e:
            print(f"Error processing table: {e}")
            continue

        prev_block = block

    return parsed_abbrs


def get_text(file):
    full_text = []
    doc = Document(file)
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)


def get_metadata(file):
    doc = Document(file)
    metadata = {}
    prop = doc.core_properties
    metadata["author"] = prop.author
    metadata["created"] = prop.created
    metadata["modified"] = prop.modified
    metadata["last_modified_by"] = prop.last_modified_by
    return metadata
