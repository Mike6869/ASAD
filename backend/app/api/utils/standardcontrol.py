import re
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import OrderedDict as odict
from typing import Dict, List, OrderedDict
from docx.shared import Pt

from models.abbreviation import get_abbreviations_by_file_id
from models.file import get_path
from models.document import get_document_text


PT_TO_CM = 0.0352778  # Константа для преобразования пунктов в сантиметры


def analyze_count_abbreviations(file_id: int) -> Dict:
    """Проверяет частоту использования аббревиатур и возвращает результат в формате JSON"""
    result = {
        "abbreviations": {
            "error": [],
            "correct": []
        },
        "error": None
    }

    text = get_document_text(file_id)
    abbrs = get_abbreviations_by_file_id(file_id)
    print(abbrs)
    try:
        # Регулярные выражения для поиска аббревиатур
        re_abbreviations = re.compile(r'\b[А-ЯA-Z]{3,}\b')
        re_abbreviation = re.compile(r'^[А-ЯA-Z]{3,}$')

        # Находим все аббревиатуры в тексте
        abbrs_in_text = re_abbreviations.findall(text)
        # Фильтруем только те, что есть в глоссарии
        filtered_abbrs = set(filter(lambda x: re_abbreviation.match(x), abbrs))

        # Считаем количество вхождений
        counter: Dict[str, int] = {}
        for abbr in filtered_abbrs:
            counter[abbr] = 0

        for abbr in abbrs_in_text:
            if abbr in counter:
                counter[abbr] += 1

        # Разделяем на ошибочные и корректные
        for abbr, count in counter.items():
            abbr_info = {
                "abbreviation": abbr,
                "count": count
            }

            if count <= 4:  # Меньше или равно 4
                result["abbreviations"]["error"].append(abbr_info)
            else:  # Больше 4
                result["abbreviations"]["correct"].append(abbr_info)

    except Exception as e:
        result["error"] = None

    return result


def get_raw_document_text(file_id: int) -> str:
    """Получает текст из DOCX с сохранением всех символов"""
    try:
        file = get_path(file_id)[0]
        doc = Document(file)
        text_parts = []

        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        return '\n'.join(text_parts)

    except Exception as e:
        raise RuntimeError(f"Ошибка чтения файла: {str(e)}")


def check_abbreviations(text: str) -> Dict[str, List[Dict]]:
    """Проверяет аббревиатуры на наличие неразрывных пробелов"""
    patterns = OrderedDict([
        ('ООО', r'(ООО)(\s)([^а-яё])'),
        ('ПАО', r'(ПАО)(\s)([^а-яё])'),
        ('АО', r'(АО)(\s)([^а-яё])'),
        ('№', r'(№)(\s)(\d)'),
        ('ГОСТ', r'(ГОСТ)(\s)(\d)')
    ])

    results = {key: [] for key in patterns.keys()}

    for name, pattern in patterns.items():
        for match in re.finditer(pattern, text):
            space = match.group(2)
            is_nbsp = space == '\xa0' or space == '&nbsp;'

            start = max(0, match.start() - 20)
            end = min(len(text), match.end() + 20)

            results[name].append({
                'position': match.start(),
                'has_nbsp': is_nbsp,
                'text_before': text[start:match.start()],
                'text_after': text[match.end():end],
                'full_match': match.group(),
                'recommendation': "Замените пробел на неразрывный" if not is_nbsp else None
            })


    structured = odict()
    for name, items in results.items():
        structured[name] = {
            'errors': [item for item in items if not item['has_nbsp']],
            'correct': [item for item in items if item['has_nbsp']]
        }
    return structured


def check_numbers(text: str) -> List[Dict]:
    """Проверяет пробелы после чисел"""
    # Паттерн для чисел (целые, десятичные, с разделителями)
    pattern = r'''
        (?P<number>
            \d{1,3}(?:[ \xa0]?\d{3})*  # Целые числа с разделителями
            (?:,\d+)?                   # Десятичная часть
        )
        (?P<space>[ \xa0])
        (?P<after>\S+)
    '''
    results = []

    for match in re.finditer(pattern, text, re.VERBOSE | re.UNICODE):
        space = match.group('space')
        is_nbsp = space == '\xa0'

        start = max(0, match.start() - 20)
        end = min(len(text), match.end() + 20)

        results.append({
            'position': match.start(),
            'number': match.group('number'),
            'after': match.group('after'),
            'has_nbsp': is_nbsp,
            'text_before': text[start:match.start()],
            'text_after': text[match.end():end],
            'full_match': match.group(),
            'recommendation': "Используйте неразрывный пробел" if not is_nbsp else None
        })

    return {
        'errors': [item for item in results if not item['has_nbsp']],
        'correct': [item for item in results if item['has_nbsp']]
    }


def check_initials(text: str) -> List[Dict]:
    """Проверяет ФИО в обоих форматах на правильные пробелы"""
    patterns = [
        {
            'name': "И.О. Фамилия",
            'regex': r'''
                (?P<before_context>.{0,10})
                (?P<before_space>[ \xa0]?)
                (?P<initials>[А-ЯЁ]\.[А-ЯЁ]\.?)
                (?P<middle_space>[ \xa0])
                (?P<lastname>[А-ЯЁ][а-яё]+)
                (?P<after_context>.{0,10})
            ''',
            'space_to_check': 'middle_space'
        },
        {
            'name': "Фамилия И.О.",
            'regex': r'''
                (?P<before_context>.{0,10})
                (?P<lastname>[А-ЯЁ][а-яё]+)
                (?P<middle_space>[ \xa0])
                (?P<initials>[А-ЯЁ]\.[А-ЯЁ]\.?)
                (?P<after_space>[ \xa0]?)
                (?P<after_context>.{0,10})
            ''',
            'space_to_check': 'middle_space'
        }
    ]

    results = []

    for pattern in patterns:
        for match in re.finditer(pattern['regex'], text, re.VERBOSE):
            middle_space = match.group('middle_space')
            is_middle_nbsp = middle_space == '\xa0'

            if pattern['name'] == "И.О. Фамилия":
                before_space = match.group('before_space')
                has_before_space = bool(before_space)
                is_before_nbsp = before_space == '\xa0' if has_before_space else True
                is_correct = is_middle_nbsp and (
                            not has_before_space or is_before_nbsp)
                full_name = f"{match.group('initials')}{middle_space}{match.group('lastname')}"
            else:
                after_space = match.group('after_space')
                has_after_space = bool(after_space)
                is_after_nbsp = after_space == '\xa0' if has_after_space else True
                is_correct = is_middle_nbsp and (
                            not has_after_space or is_after_nbsp)
                full_name = f"{match.group('lastname')}{middle_space}{match.group('initials')}"

            start = max(0, match.start() - 10)
            end = min(len(text), match.end() + 10)
            context = text[start:end]

            results.append({
                'format': pattern['name'],
                'position': match.start(),
                'full_name': full_name,
                'is_correct': is_correct,
                'context': context
            })

    return {
        'errors': [item for item in results if not item['is_correct']],
        'correct': [item for item in results if item['is_correct']]
    }


def analyze_non_breaking(file_id: int) -> Dict:
    """Анализирует документ и возвращает структурированный результат"""
    result = odict([
        ('non_breaking', odict([
            ('abbreviations', odict()),
            ('numbers', odict()),
            ('initials', odict()),
        ])),
        ('error', None)
    ])

    try:
        text = get_raw_document_text(file_id)

        # Проверка аббревиатур
        abbr_results = check_abbreviations(text)
        for category, data in abbr_results.items():
            result['non_breaking']['abbreviations'][category] = {
                'errors': [{
                    'position': item['position'],
                    'context_before': item['text_before'],
                    'context_after': item['text_after'],
                    'match': item['full_match'],
                    'recommendation': item['recommendation']
                } for item in data['errors']],
                'correct': [{
                    'position': item['position'],
                    'context': item['text_before'] + item['full_match'] + item[
                        'text_after']
                } for item in data['correct']]
            }

        # Проверка чисел
        numbers_data = check_numbers(text)
        result['non_breaking']['numbers'] = {
            'errors': [{
                'position': item['position'],
                'number': item['number'],
                'next_word': item['after'],
                'context': item['text_before'] + '>>>' + item[
                    'full_match'] + '<<<' + item['text_after'],
                'recommendation': item['recommendation']
            } for item in numbers_data['errors']],
            'correct': [{
                'position': item['position'],
                'context': item['text_before'] + item['full_match'] + item[
                    'text_after']
            } for item in numbers_data['correct']]
        }

        # Проверка ФИО
        initials_data = check_initials(text)
        result['non_breaking']['initials'] = {
            'errors': [{
                'position': item['position'],
                'format': item['format'],
                'full_name': item['full_name'],
                'context': item['context']
            } for item in initials_data['errors']],
            'correct': [{
                'position': item['position'],
                'format': item['format'],
                'full_name': item['full_name'],
                'context': item['context']
            } for item in initials_data['correct']]
        }

    except Exception as e:
        result['error'] = None

    return result

def analyze_margins(file_id: int) -> Dict:
    """Проверяет поля документа и возвращает результат в формате JSON"""
    expected_margins = {
        'left': Mm(25),
        'right': Mm(15),
        'top': Mm(20),
        'bottom': Mm(20)
    }

    result = {
        "margins": {
            "left": {"expected": 25.0, "actual": None, "status": "ERROR"},
            "right": {"expected": 15.0, "actual": None, "status": "ERROR"},
            "top": {"expected": 20.0, "actual": None, "status": "ERROR"},
            "bottom": {"expected": 20.0, "actual": None, "status": "ERROR"}
        },
        "error": None
    }

    try:
        file = get_path(file_id)[0]
        doc = Document(file)
        section = doc.sections[0]

        for margin in ['left', 'right', 'top', 'bottom']:
            actual_mm = round(getattr(section, f'{margin}_margin').mm, 1)
            expected = expected_margins[margin].mm

            result["margins"][margin]["actual"] = actual_mm
            result["margins"][margin]["status"] = "OK" if abs(
                actual_mm - expected) <= 0.5 else "ERROR"

    except Exception as e:
        result["error"] = None

    return result


def get_alignment(paragraph):
    """Определяем выравнивание с учетом всех возможных источников"""
    alignment_mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE"
    }

    if paragraph.alignment is not None:
        return alignment_mapping.get(paragraph.alignment, "UNKNOWN")

    current_style = paragraph.style
    while current_style is not None:
        if hasattr(current_style, 'paragraph_format'):
            pf = current_style.paragraph_format
            if pf.alignment is not None:
                return alignment_mapping.get(pf.alignment, "UNKNOWN")
        current_style = current_style.base_style

    return "LEFT"

def get_indents(style_or_format):
    """Получаем все виды отступов для стиля или формата"""
    return {
        'first_line': style_or_format.first_line_indent,
        'left': style_or_format.left_indent,
        'right': style_or_format.right_indent,
        'hanging': style_or_format.first_line_indent if style_or_format.first_line_indent and style_or_format.first_line_indent < Pt(0) else None
    }

def get_effective_indents(paragraph):
    """Получаем фактические отступы с учетом всех источников"""
    direct = get_indents(paragraph.paragraph_format)

    style_values = {'first_line': 0.0, 'left': 0.0, 'right': 0.0, 'hanging': 0.0}
    current_style = paragraph.style
    while current_style is not None:
        if hasattr(current_style, 'paragraph_format'):
            pf = current_style.paragraph_format
            indents = get_indents(pf)
            for key in style_values:
                if indents[key]:
                    style_value = indents[key].pt if indents[key] else 0.0
                    style_values[key] = style_value
        current_style = current_style.base_style

    effective = {}
    for key in direct:
        direct_value = direct[key].pt if direct[key] else 0.0
        style_value = style_values[key]
        effective[key] = direct_value if direct_value != 0.0 else style_value

    return effective

def analyze_indent(file_id: id):
    file = get_path(file_id)[0]
    doc = Document(file)
    indent_data = {
        "indent": {
            "correct": [],
            "error": []
        }
    }

    excluded_styles = {
        'Шаблон_Заголовок 1',
        'Список1',
        'Шаблон_текст_нумерованный список_курсив',
        'Шаблон_Заголовок 2',
        'Шаблон_Заголовок 3',
        'List Paragraph',
        'Шаблон_Таблица_заголовок',
        'Шаблон_Заголовок 4_новый',
        'Шаблон_маркированный список_комментарий'
    }

    for i, p in enumerate(doc.paragraphs, 1):
        if not p.text.strip():
            continue

        if p.style.name in excluded_styles:
            continue

        alignment = get_alignment(p)
        if alignment != "JUSTIFY":
            continue

        try:
            indents = get_effective_indents(p)
            converted = {k: round(v * PT_TO_CM, 2) for k, v in indents.items()}

            # Определение статуса
            status = "error"
            if converted['first_line'] == 1.25:
                status = "correct"
            else:
                if (abs(converted['hanging']) == 1.25 and
                    converted['left'] == 1.25 and
                    converted['right'] == 1.25):
                    status = "correct"

            entry = {
                "text": p.text.strip(),
                "first_line_indent_cm": converted['first_line'],
                "style": p.style.name
            }

            indent_data["indent"][status].append(entry)

        except Exception as e:
            print(f"Ошибка в абзаце {i}: {str(e)}")

    return indent_data