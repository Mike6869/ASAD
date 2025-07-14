from docx import Document
import re

def clear_brackets(path: str) -> str:
    """Получает текст из DOCX, удаляя служебные поля в фигурных скобках"""

    def is_placeholder(text: str) -> bool:
        """Проверяет, является ли текст служебным полем"""
        return bool(re.fullmatch(r'\{.*?\}', text.strip()))

    try:
        doc = Document(path)
        text_parts = []

        # Обработка основного текста
        for paragraph in doc.paragraphs:
            if not is_placeholder(paragraph.text):
                text_parts.append(paragraph.text)

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not is_placeholder(cell.text):
                        text_parts.append(cell.text)

        # Обработка колонтитулов
        for section in doc.sections:
            # Верхние колонтитулы
            for paragraph in section.header.paragraphs:
                if not is_placeholder(paragraph.text):
                    text_parts.append(paragraph.text)

            # Нижние колонтитулы
            for paragraph in section.footer.paragraphs:
                if not is_placeholder(paragraph.text):
                    text_parts.append(paragraph.text)

        # Дополнительно: удаляем все оставшиеся вхождения {*}
        clean_text = re.sub(r'\{.*?\}', '', '\n'.join(text_parts))

        return clean_text

    except Exception as e:
        raise RuntimeError(f"Ошибка чтения файла: {str(e)}")